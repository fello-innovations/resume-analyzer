import io
from docx import Document
from models.schemas import ParsedResume, ResumeFile
from parsers.base_parser import BaseParser


class DocxParser(BaseParser):
    """Parses .docx files using python-docx."""

    def parse(self, file: ResumeFile, content: bytes) -> ParsedResume:
        doc = Document(io.BytesIO(content))
        parts = []
        # Paragraphs (main body)
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        # Tables (skills grids, experience tables, etc.)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        return ParsedResume(file=file, text=text)
